import datetime
import os
import re
from pathlib import Path
from typing import List, Tuple

import requests
from bs4 import BeautifulSoup
from docx import Document
from dotenv import dotenv_values
from googlesearch import search
from groq import Groq

from Backend.FolderContext import get_effective_folder


PROJECT_ROOT = Path(__file__).resolve().parent.parent
ENV_PATH = PROJECT_ROOT / ".env"
env_vars = dotenv_values(ENV_PATH)

GroqAPIKey = (
    env_vars.get("GroqAPIKey")
    or env_vars.get("GROQ_API_KEY")
    or os.getenv("GroqAPIKey")
    or os.getenv("GROQ_API_KEY")
)
GroqModel = (
    env_vars.get("GroqModel")
    or env_vars.get("GROQ_MODEL")
    or os.getenv("GroqModel")
    or os.getenv("GROQ_MODEL")
    or "llama-3.3-70b-versatile"
)


def _slug(text: str, max_len: int = 64) -> str:
    clean = re.sub(r"[^a-zA-Z0-9]+", "_", (text or "").strip()).strip("_")
    return (clean[:max_len] or "research_report").lower()


def _parse_length_constraints(query: str) -> Tuple[int, int]:
    q = (query or "").lower()
    pages = 2
    lines = 120

    m_pages = re.search(r"(\d+)\s*(pages?|pg|pgs|pges)\b", q)
    if m_pages:
        pages = max(1, min(500, int(m_pages.group(1))))

    m_lines = re.search(r"(\d+)\s*(lines?|lns|ln)\b", q)
    if m_lines:
        lines = max(20, min(20000, int(m_lines.group(1))))

    return pages, lines


def _extract_topic(query: str) -> str:
    q = (query or "").strip()
    lower = q.lower()
    prefixes = [
        "research ",
        "deep research ",
        "make research report on ",
        "create research report on ",
        "generate research report on ",
        "write research report on ",
    ]
    for p in prefixes:
        if lower.startswith(p):
            return q[len(p):].strip(" :,-")
    return q


def _fetch_page_text(url: str, timeout: int = 12) -> str:
    try:
        res = requests.get(
            url,
            timeout=timeout,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
                )
            },
        )
        res.raise_for_status()
        soup = BeautifulSoup(res.text, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = " ".join(soup.get_text(separator=" ").split())
        return text[:2600]
    except Exception:
        return ""


def _collect_google_sources(topic: str, num_results: int = 8) -> List[dict]:
    sources = []
    try:
        results = list(search(topic, advanced=True, num_results=num_results))
    except Exception:
        results = []

    for r in results:
        url = getattr(r, "url", "") or getattr(r, "link", "")
        title = getattr(r, "title", "") or "Untitled"
        description = getattr(r, "description", "") or ""
        if not url:
            continue
        snippet = _fetch_page_text(url)
        sources.append(
            {
                "type": "web",
                "title": title,
                "url": url,
                "description": description,
                "snippet": snippet,
            }
        )
    return sources


def _collect_youtube_links(topic: str, num_results: int = 5) -> List[str]:
    query = f"{topic} tutorial site:youtube.com"
    links = []
    try:
        results = list(search(query, advanced=True, num_results=num_results))
    except Exception:
        results = []

    for r in results:
        url = getattr(r, "url", "") or getattr(r, "link", "")
        if "youtube.com" in (url or "") or "youtu.be" in (url or ""):
            links.append(url)
    return links[:num_results]


def _collect_arxiv_papers(topic: str, max_results: int = 6) -> List[dict]:
    url = (
        "http://export.arxiv.org/api/query?"
        f"search_query=all:{requests.utils.quote(topic)}"
        f"&start=0&max_results={max_results}"
    )
    papers = []
    try:
        res = requests.get(url, timeout=15)
        res.raise_for_status()
        soup = BeautifulSoup(res.text, "xml")
        for entry in soup.find_all("entry"):
            title = (entry.title.text or "").strip()
            summary = (entry.summary.text or "").strip()
            link = ""
            for l in entry.find_all("link"):
                href = l.get("href")
                if href and "arxiv.org/abs/" in href:
                    link = href
                    break
            papers.append(
                {
                    "source": "arXiv",
                    "title": title,
                    "url": link,
                    "summary": summary[:900],
                }
            )
    except Exception:
        pass
    return papers


def _collect_ieee_and_crossref(topic: str, max_results: int = 6) -> List[dict]:
    papers = []
    try:
        url = f"https://api.crossref.org/works?rows={max_results}&query.title={requests.utils.quote(topic)}"
        res = requests.get(url, timeout=15)
        res.raise_for_status()
        items = res.json().get("message", {}).get("items", [])
        for item in items:
            title = ""
            title_list = item.get("title") or []
            if title_list:
                title = title_list[0]
            doi = item.get("DOI", "")
            publisher = item.get("publisher", "")
            doi_url = f"https://doi.org/{doi}" if doi else ""
            papers.append(
                {
                    "source": publisher or "Crossref",
                    "title": title or "Untitled",
                    "url": doi_url,
                    "summary": "",
                }
            )
    except Exception:
        pass
    return papers


def _build_report_markdown(
    topic: str,
    pages: int,
    lines: int,
    web_sources: List[dict],
    video_links: List[str],
    papers: List[dict],
) -> str:
    client = Groq(api_key=GroqAPIKey) if GroqAPIKey else None
    target_words = max(800, pages * 450, lines * 9)

    citations = []
    for i, s in enumerate(web_sources[:10], start=1):
        citations.append(f"[{i}] {s['title']} - {s['url']}")

    citation_offset = len(citations)
    for i, p in enumerate(papers[:10], start=1):
        n = citation_offset + i
        citations.append(f"[{n}] {p.get('title','Untitled')} - {p.get('url','')}")

    context_blob = []
    for i, s in enumerate(web_sources[:8], start=1):
        snippet = s.get("snippet") or s.get("description") or ""
        context_blob.append(f"[{i}] {s['title']}\nURL: {s['url']}\nSnippet: {snippet[:1200]}")

    paper_blob = []
    for i, p in enumerate(papers[:8], start=1):
        n = citation_offset + i
        paper_blob.append(
            f"[{n}] {p.get('title','Untitled')}\n"
            f"Source: {p.get('source','Unknown')}\n"
            f"URL: {p.get('url','')}\n"
            f"Summary: {p.get('summary','')[:800]}"
        )

    if not client:
        # Fallback without LLM.
        return (
            f"# Research Report: {topic}\n\n"
            "## Executive Summary\n"
            "Unable to generate an AI-written synthesis because Groq API key is missing.\n\n"
            "## Sources Collected\n"
            + "\n".join(citations)
            + "\n\n## Tutorial Videos\n"
            + "\n".join(f"- {u}" for u in video_links)
        )

    base_context = f"""
Topic: {topic}
Approx requested length: {pages} pages OR {lines} lines.
Target words: about {target_words}.

Web Source Context:
{chr(10).join(context_blob)}

Paper Context:
{chr(10).join(paper_blob)}

Video Links:
{chr(10).join(video_links)}

References to include:
{chr(10).join(citations)}
"""

    section_specs = [
        ("## Quick Summary", "Give a simple overview in plain language anyone can understand."),
        ("## Core Concepts Explained Simply", "Explain all key concepts as if teaching a beginner, with analogies."),
        ("## Step-by-Step Implementation Guide", "Provide a detailed implementation path from zero to production."),
        ("## Tools, Libraries, and Setup", "List exact tools, install commands, and setup sequence."),
        ("## Architecture and Design Choices", "Compare options, tradeoffs, and when to choose each option."),
        ("## Code Blueprint and Pseudocode", "Give practical code skeletons and patterns."),
        ("## Testing, Validation, and Debugging", "Explain how to test and troubleshoot common failures."),
        ("## Security, Privacy, and Ethics", "Cover risks, compliance, and safe deployment practices."),
        ("## Practical Checklist", "Provide an implementation checklist with clear action items."),
        ("## FAQ and Common Mistakes", "Beginner FAQ with clear fixes and anti-patterns."),
    ]

    # Expand section depth with requested size, but keep runtime bounded.
    section_count = min(len(section_specs), max(6, min(10, pages // 3)))
    selected_sections = section_specs[:section_count]
    section_word_target = max(250, target_words // max(1, section_count))

    system_prompt = (
        "You are an expert research report writer. "
        "Write in very clear, beginner-friendly language so even a child can follow the logic. "
        "Be practical and implementation-focused. "
        "Always include citations like [1], [2] for factual claims."
    )

    assembled = [f"# Research Report: {topic}", ""]
    for section_title, section_goal in selected_sections:
        user_prompt = f"""
Write one markdown section for a larger report.

Section title: {section_title}
Section goal: {section_goal}
Target length for this section: around {section_word_target} words.

Rules:
- Keep the section deeply practical and easy to follow.
- Use numbered steps and bullet lists where helpful.
- Include citations [n] tied to provided references.
- Include short code blocks if useful.
- Do not include a References section in this part.

Shared report context:
{base_context}
"""
        completion = client.chat.completions.create(
            model=GroqModel,
            temperature=0.35,
            top_p=1,
            max_tokens=2300,
            stream=False,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        section_text = (completion.choices[0].message.content or "").strip()
        assembled.append(section_text)
        assembled.append("")

    assembled.append("## Tutorial Videos")
    if video_links:
        assembled.extend(f"- {url}" for url in video_links)
    else:
        assembled.append("- No tutorial links found for this topic.")
    assembled.append("")

    assembled.append("## Research Papers")
    if papers:
        for p in papers[:12]:
            assembled.append(f"- {p.get('title', 'Untitled')} ({p.get('source', 'Source')}): {p.get('url', '')}")
    else:
        assembled.append("- No papers found.")
    assembled.append("")

    assembled.append("## References")
    assembled.extend(citations if citations else ["- No references captured."])
    assembled.append("")

    return "\n".join(assembled).strip()


def _write_docx_from_markdown(markdown_text: str, out_path: Path) -> Path:
    doc = Document()
    lines = (markdown_text or "").splitlines()

    in_code = False
    code_lang = "code"
    code_lines = []

    for raw in lines:
        line = raw.rstrip()
        trimmed = line.strip()

        if trimmed.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = trimmed[3:].strip() or "code"
                code_lines = []
            else:
                doc.add_paragraph(f"{code_lang}", style="Intense Quote")
                doc.add_paragraph("\n".join(code_lines) or "")
                in_code = False
                code_lang = "code"
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not trimmed:
            doc.add_paragraph("")
            continue

        if trimmed.startswith("# "):
            doc.add_heading(trimmed[2:].strip(), level=1)
            continue
        if trimmed.startswith("## "):
            doc.add_heading(trimmed[3:].strip(), level=2)
            continue
        if trimmed.startswith("### "):
            doc.add_heading(trimmed[4:].strip(), level=3)
            continue

        if re.match(r"^\d+\.\s+", trimmed):
            doc.add_paragraph(trimmed, style="List Number")
            continue
        if trimmed.startswith("- "):
            doc.add_paragraph(trimmed[2:].strip(), style="List Bullet")
            continue

        doc.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def generate_research_report(query: str) -> Tuple[bool, str]:
    topic = _extract_topic(query)
    if not topic:
        return False, "Please provide a research topic."

    pages, lines = _parse_length_constraints(query)
    web_sources = _collect_google_sources(topic, num_results=8)
    videos = _collect_youtube_links(topic, num_results=6)
    arxiv = _collect_arxiv_papers(topic, max_results=6)
    crossref = _collect_ieee_and_crossref(topic, max_results=6)
    papers = arxiv + crossref

    output_root = get_effective_folder()
    report_dir = output_root / "research_reports"
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    markdown_report = _build_report_markdown(
        topic=topic,
        pages=pages,
        lines=lines,
        web_sources=web_sources,
        video_links=videos,
        papers=papers,
    )
    filename = f"{_slug(topic)}_{timestamp}.docx"
    out_path = report_dir / filename
    _write_docx_from_markdown(markdown_report, out_path)
    return True, "Research report saved to workspace."
